from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def shade_cell(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

doc = Document()
doc.core_properties.title = "Executive Summary - DeerRun SEO Audit 2026"

# Title
title = doc.add_heading('DeerRun Treadmill', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle = doc.add_heading('Comprehensive SEO & Competitive Analysis', 2)
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

info = doc.add_paragraph()
info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
info_run = info.add_run('Date: ' + datetime.now().strftime("%B %d, %Y") + '\n')
info_run.font.size = Pt(11)

doc.add_paragraph()

# Key Findings
doc.add_heading('Executive Summary - Key Findings', 1)

summary = "DeerRun Treadmill has a strong market position and solid technical foundation, but is severely underperforming in search visibility. The brand can capture a significant market opportunity by owning the 'premium without forced subscriptions' segment, but this advantage is lost because potential customers cannot find the website.\n\nSEO Assessment:\n• Technical SEO Score: 62/100\n• Content & Keywords: 45/100\n• AI-SEO Readiness: 38/100\n• Site Architecture: 72/100\n• Mobile Optimization: 78/100\n• Competitive Position: 55/100\n\nOpportunity Size: $500M+ annual addressable market in the premium-value segment ($1,400-$1,800 price range)"

doc.add_paragraph(summary)
doc.add_paragraph()

# Three Biggest Issues
doc.add_heading('Three Biggest Issues (Fix First)', 1)

issues = [
    ('1. Missing H1 Tag', 'Fundamental SEO element missing - search engines cannot identify page topic', 'Fix: Add primary H1 heading (15 minutes)', 'Impact: +15% search visibility'),
    ('2. No Meta Description', 'Directly impacts click-through rate in search results', 'Fix: Create 155-160 character meta description (30 minutes)', 'Impact: +15% CTR improvement'),
    ('3. High Script-to-Content Ratio', 'Heavy JavaScript dominates page, minimal actual content visible', 'Fix: Defer non-critical scripts (1-2 hours)', 'Impact: Improved Core Web Vitals, faster indexing'),
]

for title_txt, problem, solution, impact in issues:
    p = doc.add_paragraph()
    p.add_run(title_txt + '\n').bold = True
    p.add_run('Problem: ' + problem + '\n')
    p.add_run(solution + '\n')
    p.add_run(impact).italic = True

doc.add_paragraph()

# Market Opportunity
doc.add_heading('Market Opportunity - Competitive Gap', 1)

market_opp = "DeerRun can own the PREMIUM VALUE SEGMENT ($1,400-$1,800) by differentiating on:\n\n1. NO FORCED SUBSCRIPTIONS\n• NordicTrack: Forced iFit ($39-44/month mandatory)\n• Peloton: Forced subscription (customer rating: 2.75 stars!)\n• DeerRun: Optional AI at $9.99/month = lowest cost\n\n2. CUSTOMER SERVICE EXCELLENCE\n• Peloton: 2.75 stars (customer service is their weakness)\n• ProForm: 2.9 stars (declining brand)\n• DeerRun: Opportunity to build 24/7 support + 30-day guarantee\n\n3. UNDERSERVED DEMOGRAPHICS\n• 55+ population: 10M+ potential customers\n• Tall users (6'2\"+ that need larger deck)\n• Heavier users (300+ lbs that need durability)\n• Apartment dwellers (need quiet operation)\n\n4. APP ECOSYSTEM FLEXIBILITY\n• Works with Apple Fitness+, Nike Training Club, Peloton App\n• Not locked into proprietary ecosystem\n• Customers choose their fitness platform"

doc.add_paragraph(market_opp)
doc.add_paragraph()

# What's in This Audit
doc.add_heading('Documents Included in This Audit', 1)

doc.add_heading('1. DeerRun_SEO_Audit_Report_2026-03-24.docx', 2)
p = doc.add_paragraph('Comprehensive technical SEO analysis covering:\n')
p.add_run('✓ Technical SEO audit (robots.txt, sitemaps, page speed, mobile UX)\n')
p.add_run('✓ Content & AI-SEO analysis (E-E-A-T signals, entity clarity)\n')
p.add_run('✓ Site architecture recommendations\n')
p.add_run('✓ Keyword opportunities & content gaps\n')
p.add_run('✓ 4-phase implementation roadmap (12 weeks)\n')
p.add_run('✓ Success metrics & KPIs to monitor\n')
p.add_run('✓ Quick wins prioritized by effort/impact\n\n')
p.add_run('Use this document for: Implementation planning, developer briefings, tracking progress')

doc.add_heading('2. Competitor_Market_Analysis_2026-03-24.docx', 2)
p = doc.add_paragraph('Competitive intelligence & market analysis covering:\n')
p.add_run('✓ Detailed profiles of 7 major competitors\n')
p.add_run('✓ Market size & growth projections\n')
p.add_run('✓ Digital fitness alternatives (real competition)\n')
p.add_run('✓ Customer pain points by competitor\n')
p.add_run('✓ DeerRun market opportunity analysis\n')
p.add_run('✓ Recommended positioning strategy\n')
p.add_run('✓ Keyword clusters to dominate\n\n')
p.add_run('Use this document for: Marketing strategy, positioning, sales messaging')

doc.add_paragraph()

# Implementation Timeline
doc.add_heading('Recommended Implementation Timeline', 1)

timeline_data = [
    ('Weeks 1-2 (PHASE 1)', 'Critical Fixes', 'Add H1 + meta description, defer scripts, create robots.txt'),
    ('Weeks 3-4 (PHASE 2)', 'High-Impact Improvements', 'XML sitemap, expand content, optimize images, canonical tags'),
    ('Weeks 5-8 (PHASE 3)', 'Content Strategy', 'Launch blog, create product guides, FAQ schema, video content'),
    ('Weeks 9-12 (PHASE 4)', 'Advanced SEO', 'Internal linking, schema enhancement, link building campaign'),
]

table = doc.add_table(rows=len(timeline_data)+1, cols=3)
table.style = 'Light Grid Accent 1'

header = table.rows[0].cells
header[0].text = 'Timeline'
header[1].text = 'Phase'
header[2].text = 'Key Actions'
for cell in header:
    shade_cell(cell, 'D3D3D3')

for i, (timeline, phase, actions) in enumerate(timeline_data, 1):
    row = table.rows[i].cells
    row[0].text = timeline
    row[1].text = phase
    row[2].text = actions

doc.add_paragraph()

# Expected Results
doc.add_heading('Expected Results (6-Month Projection)', 1)

results_data = [
    ('Organic Traffic', 'Current: <5K/month', 'Target: 25-40K/month', '5-10x increase'),
    ('Keyword Rankings', 'Current: Not ranking', 'Target: Top 10-20 for 30+ keywords', 'First page visibility'),
    ('Page Speed', 'Current: Impacted by scripts', 'Target: >80 PageSpeed score', 'Better Core Web Vitals'),
    ('Conversion Rate', 'Current: Baseline', 'Target: +50% improvement', 'More qualified leads'),
    ('Search Visibility', 'Current: Missing from reviews', 'Target: Top 3 for key terms', 'Beat major competitors'),
]

table = doc.add_table(rows=len(results_data)+1, cols=4)
table.style = 'Light Grid Accent 1'

header = table.rows[0].cells
header[0].text = 'Metric'
header[1].text = 'Current State'
header[2].text = '6-Month Target'
header[3].text = 'Impact'
for cell in header:
    shade_cell(cell, 'D3D3D3')

for i, (metric, current, target, impact) in enumerate(results_data, 1):
    row = table.rows[i].cells
    row[0].text = metric
    row[1].text = current
    row[2].text = target
    row[3].text = impact

doc.add_paragraph()

# Next Steps
doc.add_heading('Immediate Next Steps', 1)

next_steps = [
    'Review DeerRun_SEO_Audit_Report (focus on PHASE 1: Critical Fixes)',
    'Assign developers to implement 5 Phase 1 recommendations (15-60 min each)',
    'Review Competitor_Market_Analysis for positioning & messaging',
    'Set up Google Search Console & Analytics monitoring',
    'Schedule weekly SEO review meetings',
    'Plan content creation resources for Phase 2-3',
    'Create comparison content vs. NordicTrack, Peloton, Sole Fitness',
]

for i, step in enumerate(next_steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Bullet')

doc.add_paragraph()

# Contact & Questions
doc.add_heading('Questions?', 1)

questions = "For questions about this audit:\n\n• SEO Implementation: Refer to DeerRun_SEO_Audit_Report (sections 7-9)\n• Competitive Positioning: Refer to Competitor_Market_Analysis\n• Technical Details: Refer to Section 2 (Technical SEO Audit Detail) in main report\n• Content Strategy: Refer to Section 8 (Content Roadmap) in main report"
doc.add_paragraph(questions)

# Save
output_path = "output/seo-audit/2026-03-24/EXECUTIVE_SUMMARY_START_HERE.docx"
doc.save(output_path)
print(f"SUCCESS: Executive summary created at {output_path}")
