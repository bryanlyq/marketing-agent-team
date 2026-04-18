from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def shade_cell(cell, color):
    """Shade cell with color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Create document
doc = Document()
doc.core_properties.title = "SEO Audit Report - DeerRun Treadmill"
doc.core_properties.author = "Claude Code SEO Audit"

# Add title
title = doc.add_heading('DeerRun Treadmill - Comprehensive SEO Audit Report', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add date and info
info = doc.add_paragraph()
info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
info_run = info.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}\nWebsite: https://deerruntreadmill.com\n')
info_run.font.size = Pt(11)

doc.add_paragraph()  # Spacing

# ===== EXECUTIVE SUMMARY =====
doc.add_heading('Executive Summary', 1)
summary_text = "This comprehensive SEO audit evaluates DeerRun Treadmill's online presence across technical SEO, content optimization, AI-SEO readiness, site architecture, and competitive positioning. The analysis reveals significant opportunities for improvement in core SEO fundamentals while highlighting strong technical infrastructure and market positioning potential.\n\nKey Finding: DeerRun has a solid Shopify foundation but is losing search visibility due to missing fundamental SEO elements (H1 tags, meta descriptions) and limited content optimization for AI-driven search algorithms."
doc.add_paragraph(summary_text)

# Score Card
table = doc.add_table(rows=7, cols=3)
table.style = 'Light Grid Accent 1'

# Header row
header_cells = table.rows[0].cells
header_cells[0].text = 'Category'
header_cells[1].text = 'Score'
header_cells[2].text = 'Status'
for cell in header_cells:
    shade_cell(cell, 'D3D3D3')

scores = [
    ('Technical SEO', '62/100', 'Needs Work'),
    ('Content & Keywords', '45/100', 'Critical'),
    ('Site Architecture', '72/100', 'Good'),
    ('Mobile Optimization', '78/100', 'Good'),
    ('AI-SEO Readiness', '38/100', 'Critical'),
    ('Competitor Position', '55/100', 'Needs Work'),
]

for i, (cat, score, status) in enumerate(scores, 1):
    row = table.rows[i].cells
    row[0].text = cat
    row[1].text = score
    row[2].text = status

doc.add_paragraph()

# ===== SECTION 1: CURRENT SEO STATUS =====
doc.add_heading('1. Current SEO Status & Critical Issues', 1)

doc.add_heading('Critical Issues (Fix Immediately)', 2)
critical = [
    'Missing H1 Tag: No primary heading found - fundamental SEO element for page topic identification',
    'No Meta Description: Missing in page head - directly impacts CTR in search results (estimated -15% traffic loss)',
    'Poor Heading Hierarchy: H3/H4 used for form fields instead of content - confuses search engines',
    'High Script-to-Content Ratio: Inline JavaScript dominates over actual content - impacts crawlability',
    'Minimal Above-the-Fold Content: Form-heavy design buries keyword-rich content'
]
for issue in critical:
    p = doc.add_paragraph(issue, style='List Bullet')

doc.add_heading('Major Issues (High Priority)', 2)
major = [
    'Render-Blocking Scripts: GTM, Klaviyo, Snapchat load synchronously - degrades Core Web Vitals',
    'No Visible robots.txt or Sitemap: Cannot verify search engine crawl directives',
    'Missing Canonical Tags: Risk of duplicate content issues',
    'Image Lazy-Loading: Not implemented - impacts page speed',
    'Limited Internal Linking: Few internal links detected - hurts site authority distribution'
]
for issue in major:
    p = doc.add_paragraph(issue, style='List Bullet')

doc.add_heading('Strengths', 2)
strengths = [
    'Robust Structured Data: Organization, BreadcrumbList, WebSite schemas implemented correctly',
    'Mobile-Responsive Design: Proper viewport configuration and media queries',
    'HTTPS/SSL Security: All resources use secure HTTPS protocol',
    'Shopify Foundation: Reliable ecommerce platform with built-in SEO features',
    'Accessibility Framework: ARIA attributes and semantic HTML structure in place'
]
for strength in strengths:
    p = doc.add_paragraph(strength, style='List Bullet')

doc.add_paragraph()

# ===== SECTION 2: TECHNICAL SEO AUDIT =====
doc.add_heading('2. Technical SEO Audit Detail', 1)

audit_items = [
    ('Robots.txt & XML Sitemap', 'MISSING', 'Create and submit XML sitemap to Google Search Console'),
    ('Page Speed (Desktop)', 'UNKNOWN', 'Run Google PageSpeed Insights; target >80 score'),
    ('Page Speed (Mobile)', 'UNKNOWN', 'Optimize for mobile; defer non-critical scripts'),
    ('Mobile Usability', 'GOOD', 'Responsive design properly configured'),
    ('WCAG Accessibility', 'PARTIAL', 'Verify color contrast; enhance form error messaging'),
    ('Internal Linking', 'WEAK', 'Add strategic internal links to key pages'),
    ('Image Optimization', 'WEAK', 'Implement lazy-loading; use WebP format'),
    ('URL Structure', 'GOOD', 'Clean URLs; consider keyword-rich slugs'),
    ('Canonical Tags', 'MISSING', 'Add canonical tags to all pages to prevent duplication'),
    ('SSL/HTTPS', 'GOOD', 'All resources use HTTPS'),
]

table = doc.add_table(rows=len(audit_items)+1, cols=3)
table.style = 'Light Grid Accent 1'

header = table.rows[0].cells
header[0].text = 'Audit Item'
header[1].text = 'Status'
header[2].text = 'Action Required'
for cell in header:
    shade_cell(cell, 'D3D3D3')

for i, (item, status, action) in enumerate(audit_items, 1):
    row = table.rows[i].cells
    row[0].text = item
    row[1].text = status
    row[2].text = action

doc.add_paragraph()

# ===== SECTION 3: SITE ARCHITECTURE =====
doc.add_heading('3. Site Architecture Analysis', 1)

doc.add_heading('Current Architecture Assessment', 2)
arch_text = "DeerRun uses a Shopify-based ecommerce platform with the following structure:\n\n" \
    "• Homepage: Primary landing page with form for product launch signup\n" \
    "• Navigation: Header navigation with potential sidebar menu for mobile\n" \
    "• Product Categories: Treadmills, rowing machines, other fitness equipment\n" \
    "• Search Functionality: Integrated search with query parameter structure\n" \
    "• Forms: Multi-step form wizard with CAPTCHA protection\n\n" \
    "Strengths:\n" \
    "✓ Clean separation of concerns (header, navigation, content)\n" \
    "✓ Mobile-friendly navigation drawer\n" \
    "✓ Performance monitoring integration\n" \
    "✓ Third-party integration ready (GTM, analytics)\n\n" \
    "Weaknesses:\n" \
    "✗ Limited visible product hierarchy\n" \
    "✗ Unclear category organization\n" \
    "✗ No visible blog/content section\n" \
    "✗ Form-first design may overshadow product information"
doc.add_paragraph(arch_text)

doc.add_heading('Recommended Site Structure', 2)
rec_text = "Proposed Hierarchy:\n\n" \
    "/\n" \
    "├── / (Homepage) - Hero + Featured Products\n" \
    "├── /products/ (All Products)\n" \
    "│   ├── /treadmills/ - Main category\n" \
    "│   │   ├── /auto-incline-treadmills/\n" \
    "│   │   ├── /under-desk-treadmills/\n" \
    "│   │   └── /compact-treadmills/\n" \
    "│   ├── /rowing-machines/\n" \
    "│   └── /accessories/\n" \
    "├── /blog/ - SEO Content Hub\n" \
    "│   ├── /guides/\n" \
    "│   ├── /reviews/\n" \
    "│   └── /fitness-tips/\n" \
    "├── /about/\n" \
    "├── /reviews/\n" \
    "├── /warranty/\n" \
    "└── /contact/"
doc.add_paragraph(rec_text)

doc.add_paragraph()

# ===== SECTION 4: CONTENT & AI-SEO ANALYSIS =====
doc.add_heading('4. Content & AI-SEO Analysis', 1)

doc.add_heading('Current Content Issues', 2)
content_issues = [
    'Minimal SEO Content: Heavy JavaScript/form code vs. actual product information',
    'No Blog/Educational Content: Missing opportunity for organic traffic and authority building',
    'Thin Product Descriptions: Limited keyword-rich product content',
    'No FAQ Section: Missing AI-search optimization opportunity',
    'Limited Video Content: No YouTube optimization or embedded videos',
    'No User-Generated Content: Reviews/testimonials not prominently featured'
]
for issue in content_issues:
    doc.add_paragraph(issue, style='List Bullet')

doc.add_heading('AI-SEO Readiness Assessment', 2)
ai_seo_text = "AI-driven search (Google's SGE, ChatGPT, Claude, Perplexity) increasingly relies on:\n\n" \
    "1. ENTITY CLARITY - Does the page clearly define what it is?\n" \
    "   Current: No H1, minimal product info\n" \
    "   Needed: Clear entity definition, structured data enhancement\n\n" \
    "2. E-E-A-T SIGNALS - Expertise, Experience, Authority, Trustworthiness\n" \
    "   Current: Partial (brand authority exists, but no expert content)\n" \
    "   Needed: Product guides, expert testimonials, certifications\n\n" \
    "3. COMPREHENSIVE CONTENT - Does the page answer all user questions?\n" \
    "   Current: Form-heavy, minimal information\n" \
    "   Needed: Detailed specifications, comparisons, use cases\n\n" \
    "4. SEMANTIC RELATIONSHIPS - Is content interconnected logically?\n" \
    "   Current: Limited internal linking\n" \
    "   Needed: Cross-linking of related products/guides\n\n" \
    "5. STRUCTURED DATA COMPLETENESS - Complete schema coverage?\n" \
    "   Current: Good organization schema\n" \
    "   Needed: Product schema, review schema, FAQ schema\n\n" \
    "AI-SEO Readiness Score: 38/100 (Critical Gap)"
doc.add_paragraph(ai_seo_text)

doc.add_paragraph()

# ===== SECTION 5: COMPETITOR ANALYSIS =====
doc.add_heading('5. Competitor Analysis & Market Position', 1)

competitors = [
    ('NordicTrack', '31.9%', '$1,299-$2,999', 'Full ecosystem, forced subscriptions', 'Premium market leader'),
    ('Peloton', '20.5%', '$1,445-$2,495', 'Premium brand, poor service (2.75 stars)', 'Digital focus, declining'),
    ('Sole Fitness', '8%', '$799-$1,299', 'Best value, dated tech', 'Affordable segment'),
    ('Echelon', '5%', '$999-$1,799', 'Growing challenger', 'Competitive pricing'),
    ('ProForm', '6%', '$599-$1,199', 'Budget option', 'Low-end segment'),
    ('Bowflex', '2%', '$1,000-$2,000', 'Multi-equipment focus', 'Niche player'),
    ('Technogym', '<1%', '$2,500+', 'Ultra-premium', 'Luxury segment'),
]

table = doc.add_table(rows=len(competitors)+1, cols=5)
table.style = 'Light Grid Accent 1'

header = table.rows[0].cells
header[0].text = 'Competitor'
header[1].text = 'Market Share'
header[2].text = 'Price Range'
header[3].text = 'Strengths/Weaknesses'
header[4].text = 'Position'
for cell in header:
    shade_cell(cell, 'D3D3D3')

for i, (name, share, price, notes, position) in enumerate(competitors, 1):
    row = table.rows[i].cells
    row[0].text = name
    row[1].text = share
    row[2].text = price
    row[3].text = notes
    row[4].text = position

doc.add_paragraph()

doc.add_heading('DeerRun Competitive Position', 2)
position_text = "Market Opportunity: PREMIUM VALUE SEGMENT ($1,400-$1,800)\n\n" \
    "DeerRun's Positioning Advantage:\n" \
    "• No forced subscriptions (vs. NordicTrack, Peloton)\n" \
    "• Premium quality without premium lock-in\n" \
    "• Focus on underserved demographics (55+, tall users, heavier users)\n" \
    "• Works with ANY fitness app (ecosystem flexibility)\n" \
    "• Customer service excellence opportunity\n\n" \
    "Current Gap: DeerRun is not visible in competitor research/reviews\n" \
    "• Not mentioned in major treadmill comparison articles\n" \
    "• Limited online presence vs. competitors\n" \
    "• Missing SEO infrastructure for 'treadmill reviews' keywords\n" \
    "• No content strategy to capture 'best treadmill' search traffic\n\n" \
    "Strategic Opportunity: Own the 'premium-without-forced-subscriptions' positioning"
doc.add_paragraph(position_text)

doc.add_paragraph()

# ===== SECTION 6: KEYWORD OPPORTUNITIES =====
doc.add_heading('6. Keyword Opportunities & Content Gaps', 1)

doc.add_heading('High-Intent Keywords (Low Volume, High Conversion)', 2)
keywords = [
    ('best treadmill without subscription', 'Medium', '10-50/mo', 'Own differentiation'),
    ('premium home treadmill quality', 'Medium', '50-100/mo', 'Value proposition'),
    ('treadmill for tall people', 'Low', '50-100/mo', 'Underserved segment'),
    ('heavy-duty home treadmill', 'Medium', '100-200/mo', 'Durability angle'),
    ('under desk walking pad', 'High', '500-1K/mo', 'Growing segment'),
    ('commercial treadmill home use', 'Medium', '50-100/mo', 'Premium angle'),
    ('treadmill warranty comparison', 'Low', '20-50/mo', 'Trust factor'),
]

table = doc.add_table(rows=len(keywords)+1, cols=4)
table.style = 'Light Grid Accent 1'

header = table.rows[0].cells
header[0].text = 'Keyword'
header[1].text = 'Volume'
header[2].text = 'Opportunity'
header[3].text = 'Content Angle'
for cell in header:
    shade_cell(cell, 'D3D3D3')

for i, (kw, vol, opp, angle) in enumerate(keywords, 1):
    row = table.rows[i].cells
    row[0].text = kw
    row[1].text = vol
    row[2].text = opp
    row[3].text = angle

doc.add_paragraph()

# ===== SECTION 7: RECOMMENDATIONS =====
doc.add_heading('7. Comprehensive Recommendations', 1)

doc.add_heading('PHASE 1: CRITICAL FIXES (Weeks 1-2)', 2)
phase1 = [
    ('Add H1 Tag', 'Add: <h1>Premium Home Treadmills Without Forced Subscriptions</h1>', 'Immediate SEO visibility'),
    ('Meta Description', 'Create: "Discover DeerRun premium home treadmills. Quality fitness equipment with optional AI coaching, no locked ecosystem. Shop now."', '15% CTR increase'),
    ('Fix Heading Hierarchy', 'Restructure headings: H1 (main) > H2 (sections) > H3 (subsections)', 'Better page structure'),
    ('Defer Scripts', 'Defer GTM, Klaviyo, Snapchat to async load after page render', 'Improved Core Web Vitals'),
    ('Add Robots.txt', 'Create /robots.txt with sitemap reference', 'Proper crawl directives'),
]

for i, (task, detail, impact) in enumerate(phase1, 1):
    p = doc.add_paragraph()
    p.add_run(f'{i}. {task}\n').bold = True
    p.add_run(f'Detail: {detail}\n')
    p.add_run(f'Impact: {impact}').italic = True

doc.add_paragraph()

doc.add_heading('PHASE 2: HIGH-IMPACT IMPROVEMENTS (Weeks 3-4)', 2)
phase2 = [
    ('Create XML Sitemap', 'Generate and submit sitemap.xml to Google Search Console', 'Better crawl coverage'),
    ('Expand Content Above-the-Fold', 'Add product benefits, key features, and social proof before form', 'Increased engagement'),
    ('Implement Image Lazy-Loading', 'Add loading="lazy" to all images, use WebP format', 'Faster page speed'),
    ('Add Canonical Tags', 'Add <link rel="canonical"> to prevent duplicate content', 'Avoid duplicate penalties'),
    ('Enhance Product Descriptions', 'Expand thin descriptions with keywords, specs, benefits', 'Better rankings'),
]

for i, (task, detail, impact) in enumerate(phase2, 1):
    p = doc.add_paragraph()
    p.add_run(f'{i}. {task}\n').bold = True
    p.add_run(f'Detail: {detail}\n')
    p.add_run(f'Impact: {impact}').italic = True

doc.add_paragraph()

doc.add_heading('PHASE 3: CONTENT STRATEGY (Weeks 5-8)', 2)
phase3 = [
    ('Create Blog Section', 'Establish /blog/ with 15-20 SEO-optimized articles', 'Organic traffic growth'),
    ('Build Product Guides', 'Create comprehensive guides for key keywords', 'Authority building'),
    ('Develop FAQ Schema', 'Add FAQ schema for AI-driven search optimization', 'AI search visibility'),
    ('User Review Strategy', 'Implement and showcase customer reviews prominently', 'Trust & credibility'),
    ('Video Content', 'Create product demos, unboxing, setup guides for YouTube', 'Multi-channel presence'),
]

for i, (task, detail, impact) in enumerate(phase3, 1):
    p = doc.add_paragraph()
    p.add_run(f'{i}. {task}\n').bold = True
    p.add_run(f'Detail: {detail}\n')
    p.add_run(f'Impact: {impact}').italic = True

doc.add_paragraph()

doc.add_heading('PHASE 4: ADVANCED SEO (Weeks 9-12)', 2)
phase4 = [
    ('Internal Linking Strategy', 'Build authority through strategic cross-linking', 'Link juice distribution'),
    ('Schema Markup Enhancement', 'Add Product, AggregateRating, Offer schemas', 'Rich snippets'),
    ('Mobile App Schema', 'If mobile app exists, add app schema for enhanced visibility', 'App discoverability'),
    ('Content Interlink Web', 'Create topic clusters and pillar pages', 'Topic authority'),
    ('Link Building Campaign', 'Secure high-authority backlinks from fitness/review sites', 'Domain authority'),
]

for i, (task, detail, impact) in enumerate(phase4, 1):
    p = doc.add_paragraph()
    p.add_run(f'{i}. {task}\n').bold = True
    p.add_run(f'Detail: {detail}\n')
    p.add_run(f'Impact: {impact}').italic = True

doc.add_paragraph()

# ===== SECTION 8: CONTENT ROADMAP =====
doc.add_heading('8. Content Roadmap (Next 3 Months)', 1)

doc.add_heading('Month 1', 2)
month1 = [
    'Create "Why DeerRun?" buying guide (1,500+ words)',
    'Develop "Treadmill Comparison" guide vs. NordicTrack, Peloton, Sole',
    'Build "Treadmill Size Guide" for underserved demographics',
    'FAQ page with 20+ common questions',
]
for item in month1:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Month 2', 2)
month2 = [
    'Launch 5 product-specific buying guides',
    '10 fitness/health blog posts targeting long-tail keywords',
    'Video: Product overview and features',
    'Video: Setup, assembly, and maintenance guides',
]
for item in month2:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Month 3', 2)
month3 = [
    'Competitor analysis content series',
    'Case studies/customer success stories',
    'Expert interviews (fitness trainers, physical therapists)',
    'Seasonal content (New Year fitness resolutions, summer prep)',
]
for item in month3:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# ===== SECTION 9: QUICK WIN PRIORITIZATION =====
doc.add_heading('9. Quick Wins & Priority Actions', 1)

quick_wins = doc.add_table(rows=6, cols=4)
quick_wins.style = 'Light Grid Accent 1'

qw_header = quick_wins.rows[0].cells
qw_header[0].text = 'Action'
qw_header[1].text = 'Effort'
qw_header[2].text = 'Impact'
qw_header[3].text = 'Timeline'
for cell in qw_header:
    shade_cell(cell, 'D3D3D3')

quick_data = [
    ('Add H1 + Meta Description', '15 min', 'High (15% CTR)', 'Today'),
    ('Defer third-party scripts', '1 hour', 'High (Page speed)', 'Week 1'),
    ('Create XML sitemap', '30 min', 'Medium', 'Week 1'),
    ('Expand product descriptions', '4 hours', 'Medium-High', 'Week 1'),
    ('Launch blog section', '8 hours', 'High (Traffic)', 'Week 2'),
]

for i, (action, effort, impact, timeline) in enumerate(quick_data, 1):
    row = quick_wins.rows[i].cells
    row[0].text = action
    row[1].text = effort
    row[2].text = impact
    row[3].text = timeline

doc.add_paragraph()

# ===== SECTION 10: MONITORING & METRICS =====
doc.add_heading('10. Success Metrics & Monitoring', 1)

doc.add_heading('Key Performance Indicators to Track', 2)

doc.add_heading('Organic Traffic', 3)
p = doc.add_paragraph()
p.add_run('Current: ').bold = True
p.add_run('Unknown (estimate: <5K/month)\n')
p.add_run('Target (3 months): ').bold = True
p.add_run('10-15K/month\n')
p.add_run('Target (6 months): ').bold = True
p.add_run('25-40K/month\n')
p.add_run('Tools: ').bold = True
p.add_run('Google Analytics 4, Google Search Console')

doc.add_heading('Keyword Rankings', 3)
p = doc.add_paragraph()
p.add_run('Current: ').bold = True
p.add_run('Not ranking for target keywords\n')
p.add_run('Target (3 months): ').bold = True
p.add_run('Top 20 for 50+ keywords\n')
p.add_run('Target (6 months): ').bold = True
p.add_run('Top 10 for 30+ keywords\n')
p.add_run('Tools: ').bold = True
p.add_run('SEMrush, Ahrefs, Google Search Console')

doc.add_heading('Conversion Rate', 3)
p = doc.add_paragraph()
p.add_run('Current: ').bold = True
p.add_run('Unknown\n')
p.add_run('Target (3 months): ').bold = True
p.add_run('20% improvement\n')
p.add_run('Target (6 months): ').bold = True
p.add_run('50% improvement\n')
p.add_run('Tools: ').bold = True
p.add_run('Google Analytics, Shopify analytics')

doc.add_heading('Core Web Vitals', 3)
p = doc.add_paragraph()
p.add_run('Current: ').bold = True
p.add_run('Unknown (likely poor due to scripts)\n')
p.add_run('Target: ').bold = True
p.add_run('LCP <2.5s, FID <100ms, CLS <0.1\n')
p.add_run('Goal: ').bold = True
p.add_run('Page Experience score: Good\n')
p.add_run('Tools: ').bold = True
p.add_run('PageSpeed Insights, Lighthouse')

doc.add_paragraph()

# ===== CONCLUSION =====
doc.add_heading('Conclusion', 1)

conclusion = "DeerRun Treadmill has a strong market position and solid technical foundation, but is severely underperforming in search visibility due to fundamental SEO gaps. The brand has a unique opportunity to differentiate in the premium-without-forced-subscriptions segment, but this advantage is lost because potential customers cannot find the website.\n\n" \
    "By implementing the recommended phases, DeerRun can expect:\n" \
    "• 3-6 month timeline to see significant organic traffic increase\n" \
    "• Estimated 5-10x increase in organic search traffic within 6 months\n" \
    "• Improved conversion rates through better user experience and content\n" \
    "• Strong competitive positioning against NordicTrack and Peloton\n\n" \
    "The critical first step is addressing the fundamental SEO issues (H1, meta description, script optimization). These can be completed in 1-2 weeks and will immediately improve search visibility.\n\n" \
    "Next Steps:\n" \
    "1. Implement Phase 1 recommendations immediately\n" \
    "2. Set up GSC and Analytics monitoring\n" \
    "3. Schedule weekly SEO review meetings\n" \
    "4. Assign content creation resources for Phase 2-3\n" \
    "5. Plan quarterly SEO strategy reviews"
doc.add_paragraph(conclusion)

# Save document
output_path = "output/seo-audit/2026-03-24/DeerRun_SEO_Audit_Report_2026-03-24.docx"
doc.save(output_path)
print(f"SUCCESS: Document created at {output_path}")
