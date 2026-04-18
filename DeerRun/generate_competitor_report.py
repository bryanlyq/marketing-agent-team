from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def shade_cell(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

doc = Document()
doc.core_properties.title = "Competitor & Market Analysis - DeerRun Treadmill"

title = doc.add_heading('Competitor & Market Analysis Report', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

info = doc.add_paragraph()
info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
info_run = info.add_run('Date: ' + datetime.now().strftime("%B %d, %Y") + '\nDeerRun Treadmill Competitive Intelligence\n')
info_run.font.size = Pt(11)

doc.add_paragraph()

# Market Overview
doc.add_heading('Market Overview', 1)
market_overview = "The global treadmill market is valued at $4.04 billion (2026) and is projected to grow to $5.41 billion by 2035 at 3.3% CAGR. The smart fitness market is growing at 26.1% annually (much faster). Under-desk treadmills are growing at 4.57% annually, representing a significant niche."
doc.add_paragraph(market_overview)
doc.add_paragraph()

# Market Dynamics
doc.add_heading('Key Market Dynamics', 1)
dynamics = [
    'Post-pandemic normalization: Market stabilizing after initial spike',
    'Subscription fatigue: Consumers avoid forced fitness subscriptions ($40-44/month)',
    'Premiumization: Wealthy consumers seek high-end equipment without lock-in',
    'Underserved segments: 55+, tall users, heavier users (10M+ potential customers)',
    'Digital-first strategy: Younger consumers prefer app-based fitness',
    'Commercial-to-home trend: Demand for commercial-grade quality at home prices',
    'Noise concerns: Urban/apartment dwellers seeking quiet equipment',
]
for dynamic in dynamics:
    doc.add_paragraph(dynamic, style='List Bullet')
doc.add_paragraph()

# Competitive Landscape
doc.add_heading('Detailed Competitive Landscape', 1)

doc.add_heading('1. NordicTrack (31.9% Market Share)', 2)
p = doc.add_paragraph()
p.add_run('Price Range: ').bold = True
p.add_run('$1,299-$2,999\n')
p.add_run('Strengths: ').bold = True
p.add_run('Complete ecosystem, incline technology, iFit subscription\n')
p.add_run('Weaknesses: ').bold = True
p.add_run('Forced subscription, poor customer service\n')
p.add_run('Rating: ').bold = True
p.add_run('3.2/5 stars')

doc.add_heading('2. Peloton (20.5% Market Share)', 2)
p = doc.add_paragraph()
p.add_run('Price Range: ').bold = True
p.add_run('$1,445-$2,495\n')
p.add_run('Strengths: ').bold = True
p.add_run('Premium brand, live classes, strong community\n')
p.add_run('Weaknesses: ').bold = True
p.add_run('WORST customer service (2.75 stars) - MAJOR DeerRun OPPORTUNITY\n')
p.add_run('Rating: ').bold = True
p.add_run('2.75/5 stars')

doc.add_heading('3. Sole Fitness (8% Market Share)', 2)
p = doc.add_paragraph()
p.add_run('Price Range: ').bold = True
p.add_run('$799-$1,299\n')
p.add_run('Strengths: ').bold = True
p.add_run('Best value, durable, no subscription\n')
p.add_run('Weaknesses: ').bold = True
p.add_run('Dated technology, limited connectivity\n')
p.add_run('Rating: ').bold = True
p.add_run('4.0/5 stars')

doc.add_heading('4. Echelon (5% Market Share)', 2)
p = doc.add_paragraph()
p.add_run('Price Range: ').bold = True
p.add_run('$999-$1,799\n')
p.add_run('Strengths: ').bold = True
p.add_run('Competitive pricing, growing community\n')
p.add_run('Rating: ').bold = True
p.add_run('3.7/5 stars')

doc.add_heading('5. ProForm (6% Market Share)', 2)
p = doc.add_paragraph()
p.add_run('Price Range: ').bold = True
p.add_run('$599-$1,199\n')
p.add_run('Weaknesses: ').bold = True
p.add_run('Low quality, poor service, declining brand\n')
p.add_run('Rating: ').bold = True
p.add_run('2.9/5 stars')

doc.add_heading('6. Bowflex (2% Market Share)', 2)
p = doc.add_paragraph()
p.add_run('Price Range: ').bold = True
p.add_run('$1,000-$2,000\n')
p.add_run('Focus: ').bold = True
p.add_run('Multi-equipment, compact designs\n')
p.add_run('Rating: ').bold = True
p.add_run('3.5/5 stars')

doc.add_heading('7. Technogym (<1% Market Share)', 2)
p = doc.add_paragraph()
p.add_run('Price Range: ').bold = True
p.add_run('$2,500+\n')
p.add_run('Strengths: ').bold = True
p.add_run('Commercial-grade quality\n')
p.add_run('Rating: ').bold = True
p.add_run('4.5/5 stars')

doc.add_paragraph()

# DeerRun Opportunity
doc.add_heading('DeerRun Market Opportunity', 1)
opportunity_text = "MARKET GAP: Premium Home Fitness Without Forced Subscriptions ($1,400-$1,800)\n\nDeerRun's Unique Position:\n✓ No forced ecosystem lock-in\n✓ Premium quality with no subscription model\n✓ Works with ANY fitness app\n✓ Optional AI coaching at $9.99/mo (vs. $40/mo)\n✓ Targets underserved: 55+, tall people, heavier users\n\nCurrent Gap:\n✗ Not visible in competitor research\n✗ No search presence\n✗ Not mentioned in review articles\n✗ Missing from fitness equipment comparisons\n\nAddressable Market:\n• 2.1M households in premium segment seeking alternatives\n• 3.8M underserved demographic\n• $500M+ annual market opportunity"
doc.add_paragraph(opportunity_text)

doc.add_paragraph()

# Customer Pain Points
doc.add_heading('Competitor Customer Pain Points', 1)
pain_points = [
    ('Subscription Fatigue', '23% of complaints', 'DeerRun: Optional at $9.99/mo'),
    ('Assembly Difficulties', '23% of complaints', 'DeerRun: White-glove setup included'),
    ('Size Limitations', '19% of complaints', 'DeerRun: Multiple sizes for all bodies'),
    ('Noise Issues', '18% of complaints', 'DeerRun: Whisper-quiet operation'),
    ('Poor Customer Service', '15% of complaints', 'DeerRun: 24/7 support guarantee'),
]

table = doc.add_table(rows=len(pain_points)+1, cols=3)
table.style = 'Light Grid Accent 1'

header = table.rows[0].cells
header[0].text = 'Pain Point'
header[1].text = 'Frequency'
header[2].text = 'DeerRun Solution'
for cell in header:
    shade_cell(cell, 'D3D3D3')

for i, (pain, freq, solution) in enumerate(pain_points, 1):
    row = table.rows[i].cells
    row[0].text = pain
    row[1].text = freq
    row[2].text = solution

doc.add_paragraph()

# Positioning
doc.add_heading('Recommended Positioning Strategy', 1)
doc.add_heading('Tagline:', 2)
doc.add_paragraph('"Premium home fitness equipment without forced subscriptions"')

positioning = "Target: Serious fitness enthusiasts aged 35-65, price-conscious premium buyers\n\nKey Differentiators:\n1. No ecosystem lock-in - Works with any fitness app\n2. Optional AI coaching at $9.99/month\n3. Lifetime warranty + 24/7 support\n4. Built for underserved: 55+, tall users, heavier users\n5. Apartment-friendly, whisper-quiet operation\n\nPricing: $1,400-$1,800 (Premium-Value Segment)\n• Above Sole but below NordicTrack/Peloton\n• No mandatory subscription = lowest total cost"
doc.add_paragraph(positioning)

doc.add_paragraph()

# Save
output_path = "output/seo-audit/2026-03-24/Competitor_Market_Analysis_2026-03-24.docx"
doc.save(output_path)
print(f"SUCCESS: Competitor analysis created at {output_path}")
